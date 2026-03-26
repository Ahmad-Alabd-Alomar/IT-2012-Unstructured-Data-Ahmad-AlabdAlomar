import os
from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas

def create_folders():
    """Creates the required Lab 4 folders."""
    folders = ['data/raw/word', 'data/raw/excel', 'data/raw/pdf']
    for folder in folders:
        os.makedirs(folder, exist_ok=True)

def create_mock_word():
    """Generates a Word doc with text and tables for Lab 4."""
    doc = Document()
    doc.add_heading('Python Course Syllabus', 0)
    
    doc.add_paragraph('This is standard text describing the Python course.')
    doc.add_paragraph('Column 1: Intro to Python\t\tColumn 2: Advanced Data Engineering') # Fake two-column
    
    # Table WITH borders (default)
    doc.add_heading('Grading Rubric (Bordered Table)', level=2)
    table1 = doc.add_table(rows=2, cols=2)
    table1.style = 'Table Grid'
    table1.cell(0, 0).text = 'Assignment'
    table1.cell(0, 1).text = 'Points'
    table1.cell(1, 0).text = 'Lab 4'
    table1.cell(1, 1).text = '100'

    # Table WITHOUT borders
    doc.add_heading('Course Schedule (No Borders)', level=2)
    table2 = doc.add_table(rows=2, cols=2)
    table2.cell(0, 0).text = 'Week 1'
    table2.cell(0, 1).text = 'APIs'
    table2.cell(1, 0).text = 'Week 2'
    table2.cell(1, 1).text = 'Document Extraction'
    
    doc.save('data/raw/word/mock_syllabus.docx')
    print("✅ Created Word Document: data/raw/word/mock_syllabus.docx")

def create_mock_excel():
    """Generates an Excel file with formulas and tables for Lab 4."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Course Pricing"
    
    ws.append(["Course Name", "Price", "Tax", "Total"])
    ws.append(["Python Basics", 50, 5, "=B2+C2"]) # Formula included!
    ws.append(["Data Engineering", 100, 10, "=B3+C3"])
    ws.append(["Total Revenue", "", "", "=SUM(D2:D3)"])
    
    wb.save('data/raw/excel/mock_finances.xlsx')
    print("✅ Created Excel Document: data/raw/excel/mock_finances.xlsx")

def create_mock_pdf():
    """Generates a basic PDF for text extraction testing."""
    c = canvas.Canvas("data/raw/pdf/mock_overview.pdf")
    c.drawString(100, 800, "Education Pipeline PDF Extraction Test")
    c.drawString(100, 780, "This is a normal text paragraph inside a PDF file.")
    c.drawString(100, 760, "We will use pdfplumber to extract this text.")
    c.save()
    print("✅ Created PDF Document: data/raw/pdf/mock_overview.pdf")

if __name__ == "__main__":
    create_folders()
    create_mock_word()
    create_mock_excel()
    create_mock_pdf()
    print("🎉 All Lab 4 mock files successfully generated!")