import os
import pdfplumber
import docx
import openpyxl
import chardet
import logging

def safe_decode(raw_bytes):
    """
    Lab 4 Requirement: Uses chardet to detect and handle encoding issues.
    Ensures extracted text is readable and correctly decoded.
    """
    if not isinstance(raw_bytes, bytes):
        return raw_bytes
        
    detected = chardet.detect(raw_bytes)
    encoding = detected['encoding'] or 'utf-8'
    try:
        return raw_bytes.decode(encoding)
    except Exception as e:
        logging.warning(f"Decoding failed with {encoding}, falling back to utf-8. Error: {e}")
        return raw_bytes.decode('utf-8', errors='replace')

def extract_from_pdf(file_path):
    """Extracts text and tables from a PDF file using pdfplumber."""
    extracted_data = {
        "text": [],
        "tables": []
    }
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # Extract Text
                text = page.extract_text()
                if text:
                    extracted_data["text"].append(text.strip())
                
                # Extract Tables
                tables = page.extract_tables()
                for table in tables:
                    # Clean out empty rows/cells
                    cleaned_table = [[cell for cell in row if cell] for row in table if any(row)]
                    if cleaned_table:
                        extracted_data["tables"].append(cleaned_table)
                        
        return extracted_data
    except Exception as e:
        logging.error(f"Error extracting PDF {file_path}: {e}")
        return None

def extract_from_word(file_path):
    """Extracts paragraphs and tables from a Word (.docx) file."""
    extracted_data = {
        "text": [],
        "tables": []
    }
    try:
        doc = docx.Document(file_path)
        
        # Extract normal text paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Pass through our chardet helper to ensure safe encoding strings
                safe_text = safe_decode(para.text.encode('utf-8'))
                extracted_data["text"].append(safe_text)
                
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            extracted_data["tables"].append(table_data)
            
        return extracted_data
    except Exception as e:
        logging.error(f"Error extracting Word doc {file_path}: {e}")
        return None

def extract_from_excel(file_path):
    """Extracts rows from all sheets in an Excel (.xlsx) file."""
    extracted_data = {
        "sheets": {}
    }
    try:
        # data_only=True evaluates the formulas into actual values
        wb = openpyxl.load_workbook(file_path, data_only=True)
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_data = []
            
            for row in sheet.iter_rows(values_only=True):
                # Filter out completely empty rows
                if any(cell is not None for cell in row):
                    # Convert tuples to lists and replace None with empty string
                    clean_row = [str(cell) if cell is not None else "" for cell in row]
                    sheet_data.append(clean_row)
                    
            extracted_data["sheets"][sheet_name] = sheet_data
            
        return extracted_data
    except Exception as e:
        logging.error(f"Error extracting Excel {file_path}: {e}")
        return None